from __future__ import annotations

import datetime as dt
import json
import zipfile
from importlib import resources
from pathlib import Path

import httpx
import pytest
from pydantic import ValidationError

import app.project_planner.domain_playbooks as domain_playbooks
import app.project_planner.postprocess as project_postprocess
from app.core.config import settings
from app.db.session import get_session
from app.llm.provider import LLMResponse
from app.main import create_app
from app.project_planner import budget_catalog
from app.project_planner.budget import (
    BUDGET_CONCEPT_COST_ALIGNMENT_WARNING,
    BUDGET_LLM_OVERWRITE_WARNING,
)
from app.project_planner.budget_catalog import (
    CATALOG_EMERGENCY_FALLBACK_WARNING,
    BudgetCatalogError,
    load_builtin_budget_catalog,
    parse_budget_catalog,
    resolve_budget_item,
    resolve_budget_items,
)
from app.project_planner.docx_export import (
    _clean_text,
    _generated_at_text,
    export_project_report_docx,
)
from app.project_planner.domain_playbooks import (
    GENERAL_PROJECT_TYPE,
    HIGH_CONFIDENCE_THRESHOLD,
    DomainPlaybookError,
    build_playbook_prompt_context,
    classify_project_type,
    load_domain_playbooks,
    parse_domain_playbooks,
    select_playbook,
)
from app.project_planner.generator import FALLBACK_VALIDATION_WARNING, generate_project_report
from app.project_planner.llm_normalizer import normalize_llm_project_report_json
from app.project_planner.mock_generator import build_mock_report
from app.project_planner.postprocess import (
    ROADMAP_DEADLINE_CORRECTION_WARNING,
    ROADMAP_DEADLINE_FALLBACK_WARNING,
    ROADMAP_PAST_DEADLINE_FALLBACK_WARNING,
    ROADMAP_SHORT_DEADLINE_WARNING,
    ROADMAP_START_DATE_CORRECTION_WARNING,
    build_gantt_from_roadmap,
    postprocess_project_report,
)
from app.project_planner.prompts import (
    PROJECT_REPORT_JSON_SKELETON,
    PROJECT_REPORT_JSON_SKELETON_TEXT,
    SYSTEM_PROMPT,
    build_user_prompt,
)
from app.project_planner.schemas import ProjectPlannerInput, ProjectReport
from app.project_planner.validators import (
    RECOMMENDED_CONCEPT_MISMATCH_WARNING,
    build_clarifications,
    validate_project_report,
)


def _payload() -> ProjectPlannerInput:
    return ProjectPlannerInput(
        idea="Провести фестиваль талантов в Уральском банке с региональным охватом",
        deadline=dt.date.today() + dt.timedelta(days=90),
        budget=1_500_000,
        geography="Свердловская область",
        stakeholders="HR, руководители направлений, сотрудники банка",
        current_resources="Команда внутренних коммуникаций и площадки банка",
        technology_constraints="Использовать только согласованные внутренние каналы",
        project_accents="Учесть 185-летие Сбера и идеи фестиваля 2023 года",
    )


def _it_service_payload() -> ProjectPlannerInput:
    return ProjectPlannerInput(
        idea=(
            "Запустить внутренний сервис регистрации инициатив с маршрутизацией заявок, "
            "пилотом и поддержкой пользователей"
        ),
        deadline=dt.date.today() + dt.timedelta(days=120),
        budget=2_000_000,
        geography="Челябинская область",
        stakeholders="Product owner, ИТ, ИБ, support/admin, руководители направлений",
        current_resources="Команда бизнес-анализа и тестовый MVP-контур",
        technology_constraints=(
            "Внутренний контур банка, интеграции, QA/testing, ограничения SaaS и ИБ"
        ),
        project_accents="Сделать пилот на одном направлении и подготовить масштабирование",
    )


def _general_payload() -> ProjectPlannerInput:
    return ProjectPlannerInput(
        idea="Подготовить управляемую инициативу для внутреннего согласования",
        deadline=dt.date.today() + dt.timedelta(days=90),
        geography="Свердловская область",
        stakeholders="Бизнес-заказчик и проектная команда",
        current_resources="Базовая команда проекта",
        project_accents="Собрать понятный MVP-план",
    )


def _report_json() -> dict:
    return build_mock_report(_payload()).model_dump(mode="json")


def _budget_catalog_json() -> dict:
    return json.loads(
        resources.files(budget_catalog.CATALOG_RESOURCE_PACKAGE)
        .joinpath(budget_catalog.CATALOG_RESOURCE_NAME)
        .read_text(encoding="utf-8")
    )


def _domain_playbooks_json() -> dict:
    return json.loads(
        resources.files(domain_playbooks.PLAYBOOK_RESOURCE_PACKAGE)
        .joinpath(domain_playbooks.PLAYBOOK_RESOURCE_NAME)
        .read_text(encoding="utf-8")
    )


def _report_json_with_roadmap_dates(payload: ProjectPlannerInput, start_date: dt.date) -> dict:
    raw = build_mock_report(payload, today=start_date - dt.timedelta(days=30)).model_dump(
        mode="json"
    )
    for phase_index, phase in enumerate(raw["roadmap"]):
        phase_start = start_date + dt.timedelta(days=phase_index * 10)
        phase_end = phase_start + dt.timedelta(days=9)
        phase["start_date"] = phase_start.isoformat()
        phase["end_date"] = phase_end.isoformat()
        for milestone_index, milestone in enumerate(phase["milestones"], start=1):
            due_date = min(phase_start + dt.timedelta(days=milestone_index * 2), phase_end)
            milestone["due_date"] = due_date.isoformat()
    return raw


class FakeBadProvider:
    def __init__(self) -> None:
        self.calls = 0

    async def generate_text(self, messages: list[dict], **kwargs) -> LLMResponse:  # noqa: ARG002
        self.calls += 1
        return LLMResponse(content="not json", model_name="fake-bad")


class FakeErrorProvider:
    def __init__(self) -> None:
        self.calls = 0

    async def generate_text(self, messages: list[dict], **kwargs) -> LLMResponse:  # noqa: ARG002
        self.calls += 1
        raise RuntimeError("fake provider is unavailable")


class FakeJsonProvider:
    def __init__(self, payload: dict) -> None:
        self.payload = payload
        self.calls = 0

    async def generate_text(self, messages: list[dict], **kwargs) -> LLMResponse:  # noqa: ARG002
        self.calls += 1
        return LLMResponse(
            content=json.dumps(self.payload, ensure_ascii=False), model_name="fake-json"
        )


def _build_test_client(db_session):
    app = create_app()

    async def override_session():
        yield db_session

    app.dependency_overrides[get_session] = override_session
    transport = httpx.ASGITransport(app=app)
    client = httpx.AsyncClient(transport=transport, base_url="http://test")
    return app, client


def _docx_text(path) -> str:
    from docx import Document

    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_project_planner_clarifications_allow_assumptions_after_default_limit():
    payload = ProjectPlannerInput(idea="коротко", questions_asked_count=3)
    response = build_clarifications(payload)
    assert response.can_generate_with_assumptions is True
    assert response.default_limit == 3


def test_mock_report_is_valid_and_contains_source_input():
    payload = _payload()
    report = build_mock_report(payload)
    assert report.source_input.project_accents == payload.project_accents
    assert len(report.roadmap) >= 4
    assert len(report.concepts) == 3
    assert not validate_project_report(report)


def test_mock_report_does_not_include_llm_budget_path_warnings():
    report = build_mock_report(_payload())

    assert BUDGET_LLM_OVERWRITE_WARNING not in report.warnings
    assert BUDGET_CONCEPT_COST_ALIGNMENT_WARNING not in report.warnings


def test_project_planner_prompt_contains_full_json_skeleton_and_short_form_bans():
    ProjectReport.model_validate(PROJECT_REPORT_JSON_SKELETON)

    assert "Полный JSON skeleton текущей ProjectReport schema" in SYSTEM_PROMPT
    assert PROJECT_REPORT_JSON_SKELETON_TEXT in SYSTEM_PROMPT
    for field_name in ProjectReport.model_fields:
        assert f'"{field_name}"' in PROJECT_REPORT_JSON_SKELETON_TEXT

    for field_name in (
        "title",
        "goal",
        "tasks",
        "target_audience",
        "success_criteria",
        "relevance_for_ural_bank",
        "financial_items",
        "financial_total",
        "material_resources",
        "information_resources",
        "activity",
        "responsible",
        "accountable",
        "consulted",
        "informed",
        "scenario_steps",
        "advantages",
        "disadvantages",
        "estimated_cost",
        "effort_level",
        "effort_factors",
        "differences",
        "bullets",
    ):
        assert f'"{field_name}"' in PROJECT_REPORT_JSON_SKELETON_TEXT

    for instruction in (
        "passport не может быть только {name, acronym}",
        "team не может быть list[str]",
        "concepts не может быть list[str]",
        "raci не может быть dict по phase names",
        "presentation_outline не может быть list[str]",
        "defense_script должен быть string, не array",
        "Все ключи обязательны",
        "Верни только JSON без комментариев вне JSON",
        "name, start_date, end_date, milestones",
        "title, due_date, description",
        "YYYY-MM-DD",
        "Roadmap and milestone dates must not be later than the user deadline",
        "final phase must end on or before deadline",
        "JSON arrays",
        "strings",
    ):
        assert instruction in SYSTEM_PROMPT


def test_domain_playbooks_load_validate_and_stay_compact():
    playbooks = load_domain_playbooks()

    assert set(playbooks) == {"it_service", "event", "general"}
    for project_type, playbook in playbooks.items():
        assert playbook.project_type == project_type
        assert len(playbook.prompt_context_summary) <= 700
        assert len(playbook.prompt_context_summary) <= 1200
        assert len(playbook.phases) >= 4
        assert len(playbook.roles) >= 3
        assert len(playbook.concept_patterns) == 3
        assert playbook.raci_defaults.responsible
        assert playbook.raci_defaults.accountable
        assert playbook.raci_defaults.consulted
        assert playbook.raci_defaults.informed


def test_domain_playbooks_validation_rejects_missing_raci_defaults():
    data = _domain_playbooks_json()
    data["playbooks"][0].pop("raci_defaults")

    with pytest.raises(DomainPlaybookError):
        parse_domain_playbooks(data)


def test_domain_playbooks_validation_rejects_malformed_raci_defaults():
    data = _domain_playbooks_json()
    data["playbooks"][0]["raci_defaults"]["consulted"] = []

    with pytest.raises(DomainPlaybookError):
        parse_domain_playbooks(data)


def test_domain_playbooks_validation_rejects_unknown_raci_role():
    data = _domain_playbooks_json()
    data["playbooks"][0]["raci_defaults"]["responsible"] = "Несуществующая роль"

    with pytest.raises(DomainPlaybookError):
        parse_domain_playbooks(data)


def test_domain_playbook_safe_fallback_keeps_mock_and_prompt_valid(monkeypatch, caplog):
    def fail_load():
        raise DomainPlaybookError("broken test playbook")

    domain_playbooks.load_domain_playbooks_safe.cache_clear()
    monkeypatch.setattr(domain_playbooks, "load_domain_playbooks", fail_load)
    caplog.set_level("WARNING")
    try:
        report = build_mock_report(_payload())
        prompt = build_user_prompt(_payload())
    finally:
        domain_playbooks.load_domain_playbooks_safe.cache_clear()

    ProjectReport.model_validate(report.model_dump())
    warnings_text = "\n".join(report.warnings)
    assert "- project_type: general" in prompt
    assert "safe general fallback" in caplog.text
    assert "Traceback" not in warnings_text
    assert "ValidationError" not in warnings_text
    assert "Field required" not in warnings_text


def test_domain_playbook_classifier_detects_it_service_scenario():
    classification = classify_project_type(_it_service_payload())
    playbook, selected_classification = select_playbook(_it_service_payload())

    controlled_keywords = set(playbook.strong_keywords) | set(playbook.support_keywords)
    assert classification.project_type == "it_service"
    assert classification.confidence >= HIGH_CONFIDENCE_THRESHOLD
    assert selected_classification == classification
    assert playbook.project_type == "it_service"
    assert set(classification.matched_keywords) <= controlled_keywords
    assert "маршрутизацией заявок" not in classification.matched_keywords


def test_domain_playbook_classifier_detects_event_scenario():
    classification = classify_project_type(_payload())

    assert classification.project_type == "event"
    assert "фестиваль" in classification.matched_keywords


def test_domain_playbook_classifier_uses_boundaries_for_short_keywords():
    flexible = classify_project_type(ProjectPlannerInput(idea="Гибкая инициатива для команды"))
    it_service = classify_project_type(ProjectPlannerInput(idea="ИБ и внутренний контур"))

    assert flexible.project_type == GENERAL_PROJECT_TYPE
    assert it_service.project_type == "it_service"
    assert "иб" in {keyword.lower() for keyword in it_service.matched_keywords}


@pytest.mark.parametrize(
    "idea",
    (
        "Сделать сервис",
        "Организовать регистрацию",
        "Сделать сервис регистрации",
        "Уточнить внутреннюю инициативу",
    ),
)
def test_domain_playbook_classifier_keeps_low_confidence_inputs_general(idea):
    classification = classify_project_type(ProjectPlannerInput(idea=idea))

    assert classification.project_type == GENERAL_PROJECT_TYPE
    assert classification.confidence == 0
    assert classification.matched_keywords == ()


def test_build_mock_report_uses_it_service_playbook_hints():
    report = build_mock_report(_it_service_payload())
    text = "\n".join(
        [
            *(role.title for role in report.team),
            *(report.resources.material_resources),
            *(report.resources.information_resources),
            *(report.passport.risks),
            *(phase.name for phase in report.roadmap),
            *(concept.name for concept in report.concepts),
        ]
    )
    normalized_text = text.lower()

    for expected in ("Специалист ИБ", "QA/testing"):
        assert expected in text
    for expected_stem in ("интеграц", "контур", "пилот"):
        assert expected_stem in normalized_text
    assert "Support/admin" in text or "support/admin" in text
    assert report.raci[0].responsible == "Бизнес-аналитик"
    assert report.raci[0].accountable == "Product owner сервиса"
    assert "Специалист ИБ" in report.raci[0].consulted
    assert not validate_project_report(report)


def test_build_mock_report_uses_event_playbook_hints():
    report = build_mock_report(_payload())
    text = "\n".join(
        [
            *(role.title for role in report.team),
            *(report.resources.material_resources),
            *(report.resources.information_resources),
            *(report.passport.risks),
            *(phase.name for phase in report.roadmap),
            *(concept.name for concept in report.concepts),
        ]
    )
    normalized_text = text.lower()

    for expected in ("Программный менеджер", "Площадка"):
        assert expected in text
    for expected_stem in ("логист", "подрядчик", "программ"):
        assert expected_stem in normalized_text
    assert "event safety" in text or "безопасность мероприятия" in text
    assert report.raci[0].responsible == "Программный менеджер"
    assert report.raci[0].accountable == "Продюсер/руководитель мероприятия"
    assert "Логистический координатор" in report.raci[0].consulted
    assert not validate_project_report(report)


def test_build_mock_report_general_playbook_remains_valid_and_non_empty():
    report = build_mock_report(_general_payload())

    assert report.team[0].title == "Руководитель проекта"
    assert report.resources.material_resources
    assert report.resources.information_resources
    assert report.passport.risks
    assert len(report.concepts) == 3
    assert report.raci[0].responsible == "Руководитель проекта"
    assert report.raci[0].accountable == "Бизнес-заказчик"
    assert not validate_project_report(report)


def test_user_prompt_includes_compact_playbook_context_without_full_json_dump():
    prompt = build_user_prompt(_it_service_payload())
    context = build_playbook_prompt_context(_it_service_payload())

    assert "Domain playbook context" in prompt
    assert "- project_type: it_service" in prompt
    assert "- confidence:" in prompt
    assert "IT-сервис: уточнить" in prompt
    assert len(context) <= 1400
    assert '"playbooks"' not in prompt
    assert "matched_keywords" not in prompt
    assert PROJECT_REPORT_JSON_SKELETON_TEXT in SYSTEM_PROMPT


def test_validate_project_report_warns_on_recommended_concept_mismatch():
    report = build_mock_report(_payload())
    report.recommended_concept.concept_name = "Несуществующая концепция"

    warnings = validate_project_report(report)

    assert RECOMMENDED_CONCEPT_MISMATCH_WARNING in warnings


@pytest.mark.parametrize(
    "concept_name",
    (
        "Концепция C — модификация портала",
        "Концепция C - модификация портала",
        "Концепция C: модификация портала",
        "Концепция C (модификация портала)",
    ),
)
def test_validate_project_report_accepts_narrow_recommended_concept_prefix(concept_name):
    report = build_mock_report(_general_payload())
    report.concepts[0].name = concept_name
    report.concepts[1].name = "Концепция A"
    report.concepts[2].name = "Концепция B"
    report.recommended_concept.concept_name = "Концепция C"

    warnings = validate_project_report(report)

    assert RECOMMENDED_CONCEPT_MISMATCH_WARNING not in warnings


def test_validate_project_report_adds_single_compact_high_confidence_domain_warning():
    report = build_mock_report(_general_payload())

    warnings = validate_project_report(report, _it_service_payload())
    domain_warnings = [
        warning for warning in warnings if "слабо отражает доменные признаки" in warning
    ]

    assert len(domain_warnings) == 1
    assert "IT-сервис" in domain_warnings[0]
    assert len(domain_warnings[0]) < 220


def test_validate_project_report_does_not_add_domain_noise_for_general_low_confidence():
    report = build_mock_report(_general_payload())

    warnings = validate_project_report(report, ProjectPlannerInput(idea="Сделать сервис"))

    assert not any("доменные признаки" in warning for warning in warnings)


def test_budget_catalog_loads_from_package_resource():
    catalog = load_builtin_budget_catalog()

    assert catalog.catalog_name == "Project Planner demo/reference budget catalog"
    assert catalog.catalog_version == "v1"
    assert catalog.default_region == "Свердловская область"
    assert {item.category_key for item in catalog.items} >= {
        "project_management",
        "contractors_expertise",
        "marketing_communications",
        "technical_support",
        "risk_reserve",
        "other",
    }


def test_budget_catalog_validation_rejects_invalid_data():
    catalog = load_builtin_budget_catalog()
    broken = {
        "catalog_name": catalog.catalog_name,
        "catalog_version": catalog.catalog_version,
        "default_region": catalog.default_region,
        "currency": "USD",
        "source_name": catalog.source_name,
        "source_date": catalog.source_date,
        "items": [],
    }

    with pytest.raises(BudgetCatalogError):
        parse_budget_catalog(broken)


@pytest.mark.parametrize(
    ("field", "value"),
    (
        ("avg_price", -1),
        ("aliases", "оборудование"),
        ("source_date", "01.06.2026"),
    ),
)
def test_budget_catalog_validation_rejects_invalid_item_values(field, value):
    data = _budget_catalog_json()
    data["items"][0][field] = value

    with pytest.raises(BudgetCatalogError):
        parse_budget_catalog(data)


@pytest.mark.parametrize(
    ("min_price", "avg_price", "max_price"),
    (
        (300_000, 280_000, 320_000),
        (200_000, 280_000, 250_000),
    ),
)
def test_budget_catalog_validation_rejects_invalid_min_avg_max(
    min_price,
    avg_price,
    max_price,
):
    data = _budget_catalog_json()
    data["items"][0]["min_price"] = min_price
    data["items"][0]["avg_price"] = avg_price
    data["items"][0]["max_price"] = max_price

    with pytest.raises(BudgetCatalogError):
        parse_budget_catalog(data)


def test_budget_catalog_validation_requires_default_row_for_standard_categories():
    data = _budget_catalog_json()
    data["items"] = [
        item
        for item in data["items"]
        if not (item["category_key"] == "contractors_expertise" and item["region"] == "default")
    ]

    with pytest.raises(BudgetCatalogError):
        parse_budget_catalog(data)


def test_budget_catalog_invalid_load_falls_back_safely(monkeypatch):
    def fail_load():
        raise BudgetCatalogError("broken test catalog")

    monkeypatch.setattr(budget_catalog, "load_builtin_budget_catalog", fail_load)

    resolution = resolve_budget_items(_payload())

    warnings_text = "\n".join(resolution.warnings)
    assert resolution.used_emergency_fallback is True
    assert resolution.financial_items
    assert CATALOG_EMERGENCY_FALLBACK_WARNING in resolution.warnings
    assert "Traceback" not in warnings_text
    assert "broken test catalog" not in warnings_text


def test_budget_resolver_uses_exact_and_default_region_lookup():
    default_payload = _payload().model_copy(update={"geography": "Свердловская область"})
    exact_payload = _payload().model_copy(update={"geography": "ХМАО"})

    default_item, default_warnings = resolve_budget_item("technical_support", default_payload)
    exact_item, exact_warnings = resolve_budget_item("technical_support", exact_payload)

    assert default_item.category == "Техническое обеспечение"
    assert default_item.amount == 260_000
    assert "регион: Свердловская область" in default_item.comment
    assert not default_warnings
    assert exact_item.category == "Техническое обеспечение"
    assert exact_item.amount == 325_000
    assert "регион: ХМАО" in exact_item.comment
    assert "метод: demo/reference regional coefficient" not in exact_item.comment
    assert not exact_warnings


def test_budget_resolver_uses_regional_coefficient_for_known_region_without_exact_row():
    payload = _payload().model_copy(update={"geography": "Челябинская область"})
    default_item, _ = resolve_budget_item(
        "project_management",
        _payload().model_copy(update={"geography": "Свердловская область"}),
    )

    item, warnings = resolve_budget_item("project_management", payload)
    expected_amount = round(
        default_item.amount
        * budget_catalog.region_coefficient("Челябинская область")
        / budget_catalog.region_coefficient("Свердловская область"),
        -3,
    )

    assert item.category == "Подготовка и управление проектом"
    assert item.amount == expected_amount
    assert "регион: Челябинская область" in item.comment
    assert "метод: demo/reference regional coefficient" in item.comment
    assert not any("регион по умолчанию" in warning for warning in warnings)


def test_budget_resolver_invalid_default_region_coefficient_uses_safe_fallback(monkeypatch):
    def fake_region_coefficient(region: str | None) -> float:
        if region == "Свердловская область":
            return 0.0
        return 0.95

    monkeypatch.setattr(budget_catalog, "region_coefficient", fake_region_coefficient)
    payload = _payload().model_copy(update={"geography": "Челябинская область"})

    resolution = resolve_budget_items(payload)

    assert resolution.used_emergency_fallback is True
    assert CATALOG_EMERGENCY_FALLBACK_WARNING in resolution.warnings


def test_budget_resolver_unknown_region_falls_back_with_warning():
    payload = _payload().model_copy(update={"geography": "Лунная база"})

    resolution = resolve_budget_items(payload)

    assert resolution.financial_items
    assert any("Регион «Лунная база» не найден" in warning for warning in resolution.warnings)
    assert all(
        "регион: Свердловская область" in item.comment for item in resolution.financial_items
    )


def test_budget_resolver_maps_alias_and_unknown_category_to_other():
    payload = _payload()

    alias_item, alias_warnings = resolve_budget_item("оборудование", payload)
    unknown_resolution = resolve_budget_items(payload, categories=("совсем неизвестная статья",))

    assert alias_item.category == "Техническое обеспечение"
    assert not alias_warnings
    assert len(unknown_resolution.financial_items) == 1
    assert unknown_resolution.financial_items[0].category == "Прочие расходы"
    assert any("Категория бюджета" in warning for warning in unknown_resolution.warnings)


def test_budget_resolver_comments_contain_compact_provenance():
    resolution = resolve_budget_items(_payload())

    assert resolution.financial_items
    for item in resolution.financial_items:
        assert "Источник:" in item.comment
        assert "каталог: Project Planner demo/reference budget catalog v1" in item.comment
        assert "дата: 2026-06-01" in item.comment
        assert "регион:" in item.comment
        assert "confidence: test" in item.comment


def test_docx_export_creates_zip_document(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "project_planner_docx_dir", str(tmp_path), raising=False)
    report = build_mock_report(_payload())
    path = export_project_report_docx(report, run_id=42)
    assert path.exists()
    assert zipfile.is_zipfile(path)


def test_docx_export_contains_demo_ready_sections(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "project_planner_docx_dir", str(tmp_path), raising=False)
    report = build_mock_report(_payload())
    path = export_project_report_docx(report, run_id=43)

    text = _docx_text(path)

    assert "Дата генерации" in text
    assert "Оценка является предварительной" in text
    assert "Целевая аудитория" in text
    assert "Риски паспорта проекта" in text
    assert "Gantt-like представление" in text
    assert "RACI" in text
    assert "Факторы трудоёмкости" in text
    assert "Сценарий защиты" in text
    assert (
        "Оценка является предварительной, сформирована по тестовым справочникам "
        "и требует экспертной проверки перед запуском проекта MVP."
    ) in text
    assert "Источник:" in text
    assert "дата: 2026-06-01" in text
    assert "confidence: test" in text
    assert " ," not in text


def test_docx_export_builds_gantt_rows_when_report_gantt_is_empty(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "project_planner_docx_dir", str(tmp_path), raising=False)
    report = build_mock_report(_payload())
    report.gantt = []

    path = export_project_report_docx(report, run_id=44)
    text = _docx_text(path)

    assert "Gantt-like представление" in text
    assert report.roadmap[0].name in text
    assert "█" in text


def test_docx_text_cleanup_removes_common_typography_artifacts():
    assert _clean_text("по тестовым , справочникам") == "по тестовым, справочникам"
    assert (
        _clean_text("по тестовым ,  справочникам .\n  Следующее  !")
        == "по тестовым, справочникам.\nСледующее!"
    )
    assert _clean_text("HR, , руководители") == "HR, руководители"
    assert (
        _clean_text("оплата площадок , оборудование призы ,")
        == "оплата площадок, оборудование призы"
    )
    assert _clean_text("HR- , команда") == "HR-команда"
    assert _clean_text("« текст » и ( значение )") == "«текст» и (значение)"
    assert _clean_text(None) == ""


def test_docx_export_cleans_messy_core_section_text(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "project_planner_docx_dir", str(tmp_path), raising=False)
    report = build_mock_report(_payload())
    report.source_input.stakeholders = "HR, , руководители направлений"
    report.source_input.current_resources = "оплата площадок , оборудование призы ,"
    report.passport.goal = "Провести Талант шоу - , без лишних пробелов ,"
    report.passport.tasks[0] = "мастер классы - , коммуникации ,"
    report.resources.material_resources[0] = "площадка , оборудование ,"
    report.raci[0].consulted = ["HR- , команда", "финансы ,"]

    path = export_project_report_docx(report, run_id=45)
    text = _docx_text(path)

    assert "HR, руководители направлений" in text
    assert "оплата площадок, оборудование призы" in text
    assert "площадка, оборудование" in text
    assert "HR-команда" in text
    assert " ," not in text


def test_frontend_docx_download_does_not_use_spa_navigation():
    repo_root = Path(__file__).resolve().parents[2]
    page_source = (repo_root / "frontend/src/pages/ProjectPlannerPage.tsx").read_text(
        encoding="utf-8"
    )
    api_source = (repo_root / "frontend/src/api/projectPlanner.ts").read_text(encoding="utf-8")
    app_source = (repo_root / "frontend/src/App.tsx").read_text(encoding="utf-8")

    assert "href={`/api/project-planner/runs/${" not in page_source
    assert 'href="/api/project-planner/runs/' not in page_source
    assert "window.location" not in api_source
    assert "window.open" not in api_source
    assert "navigate(" not in api_source

    assert "event.preventDefault();" in page_source
    assert "event.stopPropagation();" in page_source
    assert 'type="button"' in page_source

    assert "fetch(`/api/project-planner/runs/${runId}/docx`)" in api_source
    assert "URL.createObjectURL" in api_source
    assert 'link.dataset.noSpa = "true";' in api_source
    assert "link.click();" in api_source
    assert "URL.revokeObjectURL" in api_source

    assert "if (!event.isTrusted) return;" in app_source
    assert 'link.hasAttribute("download")' in app_source
    assert 'link.hasAttribute("data-no-spa")' in app_source
    assert 'link.hasAttribute("data-router-ignore")' in app_source
    assert 'href.startsWith("/api/")' in app_source
    assert 'href.startsWith("blob:")' in app_source


@pytest.mark.parametrize("tz_value", [None, "", object(), "No/Such_Timezone"])
def test_docx_generated_at_falls_back_for_invalid_timezone(tz_value, monkeypatch):
    monkeypatch.setattr(settings, "tz", tz_value, raising=False)

    generated_at = _generated_at_text()

    assert generated_at
    assert len(generated_at) == 16


def test_validate_project_report_warns_on_milestone_count_outside_required_range():
    report = build_mock_report(_payload())
    report.roadmap[0].milestones = report.roadmap[0].milestones[:2]
    report.roadmap[1].milestones = report.roadmap[1].milestones * 3

    warnings = validate_project_report(report)

    assert any("меньше 3 контрольных точек" in warning for warning in warnings)
    assert any("больше 6 контрольных точек" in warning for warning in warnings)


def test_validate_project_report_warns_on_deadline_overflow_and_empty_gantt():
    deadline = dt.date(2026, 9, 1)
    payload = _payload().model_copy(update={"deadline": deadline})
    report = build_mock_report(payload, today=dt.date(2026, 7, 1))
    report.roadmap[0].end_date = deadline + dt.timedelta(days=1)
    report.roadmap[0].milestones[0].due_date = deadline + dt.timedelta(days=1)
    report.gantt = []

    warnings = validate_project_report(report, payload)

    assert any("выходит за пользовательский дедлайн" in warning for warning in warnings)
    assert any("Gantt-like представление не заполнено" in warning for warning in warnings)


def test_validate_project_report_warns_on_roadmap_dates_before_current_date():
    current_date = dt.date(2026, 6, 1)
    payload = _payload().model_copy(update={"deadline": dt.date(2026, 9, 1)})
    report = build_mock_report(payload, today=current_date)
    report.roadmap[0].start_date = current_date - dt.timedelta(days=1)
    report.roadmap[0].milestones[0].due_date = current_date - dt.timedelta(days=1)

    warnings = validate_project_report(report, payload, current_date=current_date)

    assert any("начинается раньше даты генерации" in warning for warning in warnings)
    assert any("назначена раньше даты генерации" in warning for warning in warnings)


def test_postprocess_rebuilds_empty_gantt_from_roadmap():
    report = build_mock_report(_payload())
    report.gantt = []

    processed, fallback_warning = postprocess_project_report(report, _payload())

    assert fallback_warning is None
    assert len(processed.gantt) == len(report.roadmap)
    assert all(row.phase and row.period and "█" in row.timeline for row in processed.gantt)


def test_build_gantt_from_roadmap_uses_simple_stable_text_bars():
    report = build_mock_report(_payload())

    rows = build_gantt_from_roadmap(report.roadmap)

    assert len(rows) == len(report.roadmap)
    assert all(row.phase and row.period and set(row.timeline) <= {"█", "░"} for row in rows)


def test_llm_normalizer_converts_source_input_lists_to_strings():
    raw = _report_json()
    raw["source_input"]["stakeholders"] = ["HR", "руководители направлений"]
    raw["source_input"]["current_resources"] = ["команда коммуникаций", "площадки банка"]

    normalized = normalize_llm_project_report_json(raw, _payload())
    report = ProjectReport.model_validate(normalized)

    assert report.source_input.stakeholders == "HR; руководители направлений"
    assert report.source_input.current_resources == "команда коммуникаций; площадки банка"


def test_llm_normalizer_converts_role_and_raci_strings_to_lists():
    raw = _report_json()
    raw["team"][0]["competencies"] = "управление проектом, коммуникации"
    raw["raci"][0]["consulted"] = "HR; Финансы"
    raw["raci"][0]["informed"] = "Команда проекта"

    normalized = normalize_llm_project_report_json(raw, _payload())
    report = ProjectReport.model_validate(normalized)

    assert report.team[0].competencies == ["управление проектом", "коммуникации"]
    assert report.raci[0].consulted == ["HR", "Финансы"]
    assert report.raci[0].informed == ["Команда проекта"]


def test_llm_normalizer_repairs_roadmap_aliases_and_milestone_dates():
    raw = _report_json()
    phase = raw["roadmap"][0]
    phase["title"] = phase.pop("name")
    phase.pop("start_date")
    phase.pop("end_date")
    phase["control_points"] = phase.pop("milestones")
    milestone = phase["control_points"][0]
    due_date = milestone.pop("due_date")
    title = milestone.pop("title")
    milestone.pop("description")
    milestone["name"] = f"{title} до {due_date}"

    normalized = normalize_llm_project_report_json(raw, _payload())
    report = ProjectReport.model_validate(normalized)

    assert report.roadmap[0].name == phase["title"]
    assert report.roadmap[0].milestones[0].due_date.isoformat() == due_date
    assert report.roadmap[0].milestones[0].description == f"{title} до {due_date}"
    assert report.roadmap[0].start_date <= report.roadmap[0].end_date


def test_llm_normalizer_replaces_malformed_resources_with_calculated_structure():
    raw = _report_json()
    raw["resources"] = {
        "summary": "модель вернула свободное описание вместо сметы",
        "material_resources": "площадка; оборудование",
    }

    normalized = normalize_llm_project_report_json(raw, _payload())
    report = ProjectReport.model_validate(normalized)

    assert report.resources.financial_items
    assert report.resources.financial_total == sum(
        item.amount for item in report.resources.financial_items
    )
    assert report.resources.material_resources == ["площадка", "оборудование"]


def test_llm_normalizer_does_not_invent_missing_core_sections():
    raw = _report_json()
    raw.pop("team")
    raw["roadmap"][0] = {
        "title": "Фаза без безопасных дат",
        "control_points": [{"name": "Контрольная точка без даты"}],
    }

    normalized = normalize_llm_project_report_json(raw, _payload())

    assert "team" not in normalized
    assert "start_date" not in normalized["roadmap"][0]
    assert "end_date" not in normalized["roadmap"][0]
    with pytest.raises(ValidationError):
        ProjectReport.model_validate(normalized)


async def test_generator_normalizes_common_gigachat_shape_errors(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 2, raising=False)
    raw = _report_json()
    raw["source_input"]["stakeholders"] = ["HR", "руководители"]
    raw["source_input"]["current_resources"] = ["команда", "площадки"]
    raw["team"][0]["competencies"] = "управление проектом, коммуникации"
    raw["raci"][0]["consulted"] = "HR; Финансы"
    raw["raci"][0]["informed"] = "Команда проекта"
    raw["concepts"][0]["differences"] = ["управляемый объём", "быстрый старт"]
    raw["resources"] = {"summary": "неструктурированная ресурсная оценка"}
    phase = raw["roadmap"][0]
    phase["title"] = phase.pop("name")
    phase.pop("start_date")
    phase.pop("end_date")
    phase["control_points"] = phase.pop("milestones")
    milestone = phase["control_points"][0]
    due_date = milestone.pop("due_date")
    title = milestone.pop("title")
    milestone.pop("description")
    milestone["name"] = f"{title} {due_date}"
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(
        _payload(),
        provider=provider,
    )

    assert used_fallback is False
    assert model_name == "fake-json"
    assert provider.calls == 1
    assert report.source_input.stakeholders == "HR; руководители"
    assert report.team[0].competencies == ["управление проектом", "коммуникации"]
    assert report.raci[0].consulted == ["HR", "Финансы"]
    assert report.concepts[0].differences == "управляемый объём; быстрый старт"
    assert report.roadmap[0].name == phase["title"]
    assert report.roadmap[0].milestones[0].due_date.isoformat() == due_date
    assert report.resources.financial_items


async def test_generator_overwrites_llm_financial_items_with_backend_resolver(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 0, raising=False)
    raw = _report_json()
    raw["concepts"][0]["estimated_cost"] = 2_500_000
    raw["concepts"][1]["estimated_cost"] = 3_000_000
    raw["concepts"][2]["estimated_cost"] = 2_000_000
    raw["concepts"][2]["effort_level"] = "очень высокая"
    raw["resources"]["financial_items"] = [
        {
            "category": "LLM invented category",
            "amount": 1,
            "comment": "LLM generated this value",
        }
    ]
    raw["resources"]["financial_total"] = 1
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(_payload(), provider=provider)

    warnings_text = "\n".join(report.warnings)
    assert used_fallback is False
    assert model_name == "fake-json"
    assert provider.calls == 1
    assert all(
        item.category != "LLM invented category" for item in report.resources.financial_items
    )
    assert report.resources.financial_total == sum(
        item.amount for item in report.resources.financial_items
    )
    assert report.warnings.count(BUDGET_LLM_OVERWRITE_WARNING) == 1
    assert report.warnings.count(BUDGET_CONCEPT_COST_ALIGNMENT_WARNING) == 1
    assert [concept.estimated_cost for concept in report.concepts] == [
        round(report.resources.financial_total * 1.00, -3),
        round(report.resources.financial_total * 1.15, -3),
        round(report.resources.financial_total * 1.15, -3),
    ]
    assert all("Источник:" in item.comment for item in report.resources.financial_items)
    assert all("дата: 2026-06-01" in item.comment for item in report.resources.financial_items)
    assert all("confidence: test" in item.comment for item in report.resources.financial_items)
    assert "Traceback" not in warnings_text
    assert "ValidationError" not in warnings_text
    assert "Field required" not in warnings_text


async def test_generator_corrects_llm_roadmap_after_user_deadline(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 0, raising=False)
    monkeypatch.setattr(project_postprocess, "_current_date", lambda: dt.date(2026, 6, 1))
    deadline = dt.date(2026, 9, 1)
    payload = _payload().model_copy(update={"deadline": deadline})
    raw = _report_json_with_roadmap_dates(payload, dt.date(2026, 8, 1))
    raw["gantt"] = []
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(payload, provider=provider)

    warnings_text = "\n".join(report.warnings)
    assert used_fallback is False
    assert model_name == "fake-json"
    assert provider.calls == 1
    assert ROADMAP_DEADLINE_CORRECTION_WARNING in report.warnings
    assert all(phase.end_date <= deadline for phase in report.roadmap)
    assert all(
        milestone.due_date <= deadline for phase in report.roadmap for milestone in phase.milestones
    )
    assert report.gantt
    assert all(row.phase and row.period and "█" in row.timeline for row in report.gantt)
    assert "ValidationError" not in warnings_text
    assert "Traceback" not in warnings_text
    assert "Field required" not in warnings_text


async def test_generator_corrects_llm_roadmap_starting_before_current_date(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 0, raising=False)
    current_date = dt.date(2026, 6, 1)
    monkeypatch.setattr(project_postprocess, "_current_date", lambda: current_date)
    deadline = dt.date(2026, 9, 1)
    payload = _payload().model_copy(update={"deadline": deadline})
    raw = _report_json_with_roadmap_dates(payload, dt.date(2026, 4, 1))
    raw["gantt"] = [{"phase": "", "period": "", "timeline": ""}]
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(payload, provider=provider)

    warnings_text = "\n".join(report.warnings)
    assert used_fallback is False
    assert model_name == "fake-json"
    assert provider.calls == 1
    assert ROADMAP_START_DATE_CORRECTION_WARNING in report.warnings
    assert all(phase.start_date >= current_date for phase in report.roadmap)
    assert all(phase.end_date <= deadline for phase in report.roadmap)
    assert all(
        milestone.due_date >= current_date
        for phase in report.roadmap
        for milestone in phase.milestones
    )
    assert all(
        milestone.due_date <= deadline for phase in report.roadmap for milestone in phase.milestones
    )
    assert report.gantt
    assert all(row.phase and row.period and "█" in row.timeline for row in report.gantt)
    assert "ValidationError" not in warnings_text
    assert "Traceback" not in warnings_text
    assert "Field required" not in warnings_text


async def test_generator_marks_too_short_compressed_deadline(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 0, raising=False)
    monkeypatch.setattr(project_postprocess, "_current_date", lambda: dt.date(2026, 6, 1))
    deadline = dt.date(2026, 8, 3)
    payload = _payload().model_copy(update={"deadline": deadline})
    raw = _report_json_with_roadmap_dates(payload, dt.date(2026, 8, 1))
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(payload, provider=provider)

    assert used_fallback is False
    assert model_name == "fake-json"
    assert ROADMAP_DEADLINE_CORRECTION_WARNING in report.warnings
    assert ROADMAP_SHORT_DEADLINE_WARNING in report.warnings
    assert all(phase.end_date <= deadline for phase in report.roadmap)


async def test_generator_uses_fallback_when_current_date_is_after_deadline(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 0, raising=False)
    current_date = dt.date(2026, 6, 2)
    monkeypatch.setattr(project_postprocess, "_current_date", lambda: current_date)
    deadline = dt.date(2026, 6, 1)
    payload = _payload().model_copy(update={"deadline": deadline})
    raw = _report_json_with_roadmap_dates(payload, dt.date(2026, 5, 1))
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(payload, provider=provider)

    warnings_text = "\n".join(report.warnings)
    assert used_fallback is True
    assert model_name == "fallback"
    assert provider.calls == 1
    assert ROADMAP_PAST_DEADLINE_FALLBACK_WARNING in report.warnings
    assert BUDGET_CONCEPT_COST_ALIGNMENT_WARNING not in report.warnings
    assert "ValidationError" not in warnings_text
    assert "Traceback" not in warnings_text
    assert "Field required" not in warnings_text


async def test_generator_uses_fallback_when_roadmap_start_is_after_deadline(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 0, raising=False)
    monkeypatch.setattr(project_postprocess, "_current_date", lambda: dt.date(2026, 6, 1))
    deadline = dt.date(2026, 9, 1)
    payload = _payload().model_copy(update={"deadline": deadline})
    raw = _report_json_with_roadmap_dates(payload, deadline + dt.timedelta(days=1))
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(payload, provider=provider)

    warnings_text = "\n".join(report.warnings)
    assert used_fallback is True
    assert model_name == "fallback"
    assert provider.calls == 1
    assert ROADMAP_DEADLINE_FALLBACK_WARNING in report.warnings
    assert BUDGET_CONCEPT_COST_ALIGNMENT_WARNING not in report.warnings
    assert "ValidationError" not in warnings_text
    assert "Traceback" not in warnings_text
    assert "Field required" not in warnings_text


async def test_generator_falls_back_on_non_normalizable_json_without_traceback(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 1, raising=False)
    raw = _report_json()
    raw.pop("team")
    provider = FakeJsonProvider(raw)

    report, model_name, used_fallback = await generate_project_report(
        _payload(),
        provider=provider,
    )

    warnings_text = "\n".join(report.warnings)
    assert used_fallback is True
    assert model_name == "fallback"
    assert provider.calls == 2
    assert FALLBACK_VALIDATION_WARNING in report.warnings
    assert BUDGET_CONCEPT_COST_ALIGNMENT_WARNING not in report.warnings
    assert "ValidationError" not in warnings_text
    assert "Field required" not in warnings_text
    assert "team" not in warnings_text


async def test_generator_falls_back_on_invalid_fake_llm(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_retry_count", 1, raising=False)
    provider = FakeBadProvider()
    report, model_name, used_fallback = await generate_project_report(
        _payload(),
        provider=provider,
    )
    assert used_fallback is True
    assert model_name == "fallback"
    assert provider.calls == 2
    assert report.passport.title
    warnings_text = "\n".join(report.warnings)
    assert FALLBACK_VALIDATION_WARNING in report.warnings
    assert "ValidationError" not in warnings_text
    assert "Field required" not in warnings_text
    assert "source_input.stakeholders" not in warnings_text
    assert BUDGET_CONCEPT_COST_ALIGNMENT_WARNING not in report.warnings


async def test_generator_falls_back_without_credentials_when_mock_disabled(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    monkeypatch.setattr(settings, "gigachat_credentials", None, raising=False)

    report, model_name, used_fallback = await generate_project_report(_payload())

    assert used_fallback is True
    assert model_name == "fallback"
    assert report.passport.title
    assert any("GigaChat не настроен" in warning for warning in report.warnings)
    assert BUDGET_CONCEPT_COST_ALIGNMENT_WARNING not in report.warnings


async def test_generator_falls_back_on_provider_error_without_retrying_network(monkeypatch):
    monkeypatch.setattr(settings, "project_planner_use_mock_llm", False, raising=False)
    provider = FakeErrorProvider()

    report, model_name, used_fallback = await generate_project_report(
        _payload(),
        provider=provider,
    )

    assert used_fallback is True
    assert model_name == "fallback"
    assert provider.calls == 1
    assert report.passport.title
    warnings_text = "\n".join(report.warnings)
    assert FALLBACK_VALIDATION_WARNING in report.warnings
    assert BUDGET_CONCEPT_COST_ALIGNMENT_WARNING not in report.warnings
    assert "fake provider is unavailable" not in warnings_text


async def test_gigachat_provider_uses_env_model_with_fake_httpx(monkeypatch):
    import app.llm.gigachat_provider as gigachat_provider
    from app.llm.gigachat_provider import GigaChatLLMProvider

    class FakeResponse:
        def __init__(self, payload: dict) -> None:
            self._payload = payload

        def raise_for_status(self) -> None:
            return None

        def json(self) -> dict:
            return self._payload

    class FakeAsyncClient:
        calls: list[dict] = []

        def __init__(self, **kwargs) -> None:  # noqa: ARG002
            pass

        async def __aenter__(self):
            return self

        async def __aexit__(self, exc_type, exc, tb) -> None:  # noqa: ANN001
            return None

        async def post(self, url: str, headers=None, data=None, json=None):  # noqa: ANN001
            self.calls.append({"url": url, "headers": headers, "data": data, "json": json})
            if url == settings.gigachat_oauth_url:
                return FakeResponse({"access_token": "fake-token", "expires_at": 4_102_444_800_000})
            return FakeResponse({"choices": [{"message": {"content": '{"ok": true}'}}]})

    monkeypatch.setattr(settings, "gigachat_credentials", "fake-credentials", raising=False)
    monkeypatch.setattr(settings, "gigachat_model", "FakeGigaChatModel", raising=False)
    monkeypatch.setattr(gigachat_provider.httpx, "AsyncClient", FakeAsyncClient)

    provider = GigaChatLLMProvider()
    response = await provider.generate_text([{"role": "user", "content": "ping"}])

    assert response.content == '{"ok": true}'
    assert response.model_name == "FakeGigaChatModel"
    assert FakeAsyncClient.calls[0]["headers"]["Authorization"] == "Basic fake-credentials"
    assert FakeAsyncClient.calls[1]["headers"]["Authorization"] == "Bearer fake-token"
    assert FakeAsyncClient.calls[1]["json"]["model"] == "FakeGigaChatModel"


async def test_project_planner_api_respects_generate_with_assumptions_flag(db_session):
    app, client = _build_test_client(db_session)
    incomplete_payload = ProjectPlannerInput(idea="коротко").model_dump(mode="json")

    async with client:
        blocked_resp = await client.post(
            "/api/project-planner/runs",
            json={"input": incomplete_payload, "generate_with_assumptions": False},
        )
        empty_list_resp = await client.get("/api/project-planner/runs")
        allowed_resp = await client.post(
            "/api/project-planner/runs",
            json={"input": incomplete_payload, "generate_with_assumptions": True},
        )
    app.dependency_overrides.clear()

    assert blocked_resp.status_code == 422
    assert "Недостаточно исходных данных" in blocked_resp.json()["detail"]
    assert empty_list_resp.status_code == 200
    assert empty_list_resp.json() == []
    assert allowed_resp.status_code == 200
    allowed_run = allowed_resp.json()["run"]
    assert allowed_run["has_docx"] is True
    assert allowed_run["assumptions"]


async def test_project_planner_api_create_list_detail_and_docx(db_session):
    app, client = _build_test_client(db_session)
    payload = _payload().model_dump(mode="json")

    async with client:
        clarify_resp = await client.post("/api/project-planner/clarifications", json=payload)
        create_resp = await client.post(
            "/api/project-planner/runs",
            json={"input": payload, "generate_with_assumptions": False},
        )
        list_resp = await client.get("/api/project-planner/runs")
    app.dependency_overrides.clear()

    assert clarify_resp.status_code == 200
    assert create_resp.status_code == 200
    assert list_resp.status_code == 200

    run = create_resp.json()["run"]
    assert run["report"]["source_input"]["project_accents"] == payload["project_accents"]
    assert run["has_docx"] is True
    assert all("file_path" not in artifact for artifact in run["artifacts"])
    assert list_resp.json()[0]["id"] == run["id"]

    app, client = _build_test_client(db_session)
    async with client:
        detail_resp = await client.get(f"/api/project-planner/runs/{run['id']}")
        docx_resp = await client.get(f"/api/project-planner/runs/{run['id']}/docx")
    app.dependency_overrides.clear()

    assert detail_resp.status_code == 200
    assert detail_resp.json()["id"] == run["id"]
    assert all("file_path" not in artifact for artifact in detail_resp.json()["artifacts"])
    assert docx_resp.status_code == 200
    assert docx_resp.content.startswith(b"PK")
