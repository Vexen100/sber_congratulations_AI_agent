from __future__ import annotations

import datetime as dt
import html
import logging
import re
import zipfile
from pathlib import Path
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from app.core.config import settings
from app.project_planner.postprocess import build_gantt_from_roadmap, is_gantt_valid
from app.project_planner.schemas import ProjectReport

logger = logging.getLogger(__name__)

PRELIMINARY_NOTICE = (
    "Оценка является предварительной, сформирована по тестовым справочникам "
    "и требует экспертной проверки перед запуском проекта MVP."
)


def _safe_filename(value: str) -> str:
    clean = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in value)
    return clean.strip("_")[:80] or "project_report"


def _clean_text(value: object | None) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"([A-Za-zА-Яа-яЁё0-9])-+\s*,\s*([A-Za-zА-Яа-яЁё0-9])", r"\1-\2", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"([«(\[{])\s+", r"\1", text)
    text = re.sub(r"\s+([»)\]}])", r"\1", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"(?:,\s*){2,}", ", ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        line = re.sub(r",+$", "", line).strip()
        lines.append(line)
    return "\n".join(lines).strip()


def _paragraph_xml(text: str) -> str:
    return f"<w:p><w:r><w:t>{html.escape(_clean_text(text))}</w:t></w:r></w:p>"


def _add_heading(document, text: str, *, level: int) -> None:
    document.add_heading(_clean_text(text), level=level)


def _add_paragraph(document, text: object | None, style: str | None = None) -> None:
    if style:
        document.add_paragraph(_clean_text(text), style=style)
    else:
        document.add_paragraph(_clean_text(text))


def _format_date(value: dt.date | None) -> str:
    if value is None:
        return "не указан"
    return value.strftime("%d.%m.%Y")


def _format_money(value: float | int | None) -> str:
    if value is None:
        return "не указано"
    return f"{float(value):,.0f}".replace(",", " ") + " ₽"


def _generated_at_text() -> str:
    tz_name = getattr(settings, "tz", None)
    if not isinstance(tz_name, str) or not tz_name.strip():
        tz = dt.timezone.utc
        return dt.datetime.now(tz).strftime("%d.%m.%Y %H:%M")
    try:
        tz = ZoneInfo(tz_name.strip())
    except (ZoneInfoNotFoundError, ValueError, TypeError):
        tz = dt.timezone.utc
    return dt.datetime.now(tz).strftime("%d.%m.%Y %H:%M")


def _write_minimal_docx(report: ProjectReport, path: Path) -> None:
    gantt_rows = report.gantt if is_gantt_valid(report.gantt) else build_gantt_from_roadmap(report.roadmap)
    paragraphs = [
        "Проектный отчёт",
        report.passport.title,
        f"Дата генерации: {_generated_at_text()}",
        PRELIMINARY_NOTICE,
        "Исходные данные",
        f"Идея: {report.source_input.idea}",
        f"Дедлайн: {_format_date(report.source_input.deadline)}",
        f"Бюджет: {_format_money(report.source_input.budget)}",
        f"География: {report.source_input.geography or 'не указана'}",
        f"Акценты: {report.source_input.project_accents or 'не указаны'}",
        "Предупреждения",
        *report.warnings,
        "Допущения",
        *report.assumptions,
        "Паспорт проекта",
        report.passport.goal,
        "Дорожная карта",
        *[
            f"{phase.name}: {phase.start_date.isoformat()} - {phase.end_date.isoformat()}"
            for phase in report.roadmap
        ],
        "Gantt-like представление",
        *[f"{row.phase}: {row.period} | {row.timeline}" for row in gantt_rows],
        "Концепции",
        *[f"{concept.name}: {concept.key_idea}" for concept in report.concepts],
        "Рекомендованная концепция",
        f"{report.recommended_concept.concept_name}: {report.recommended_concept.rationale}",
    ]
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + "".join(_paragraph_xml(item) for item in paragraphs)
        + '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/></w:sectPr></w:body></w:document>'
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>"
            ),
        )
        docx.writestr(
            "_rels/.rels",
            (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
                'Target="word/document.xml"/></Relationships>'
            ),
        )
        docx.writestr("word/document.xml", document_xml)


def _add_bullets(document, title: str, items: list[str]) -> None:
    _add_heading(document, title, level=2)
    if not items:
        _add_paragraph(document, "Нет данных.")
        return
    for item in items:
        _add_paragraph(document, item, style="List Bullet")


def _add_table(document, headers: tuple[str, ...], rows: list[tuple[str, ...]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, headers, strict=True):
        cell.text = _clean_text(title)
    for values in rows:
        row = table.add_row().cells
        for cell, value in zip(row, values, strict=True):
            cell.text = _clean_text(value)
    return table


def export_project_report_docx(report: ProjectReport, *, run_id: int) -> Path:
    out_dir = Path(settings.project_planner_docx_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"project_planner_run_{run_id}_{_safe_filename(report.passport.title)}.docx"
    path = out_dir / filename

    try:
        from docx import Document

        document = Document()
        _add_heading(document, "Проектный отчёт", level=0)
        _add_paragraph(document, report.passport.title)
        _add_paragraph(document, f"Дата генерации: {_generated_at_text()}")
        _add_paragraph(document, PRELIMINARY_NOTICE)

        _add_heading(document, "Исходные данные", level=1)
        source_rows = [
            ("Идея", report.source_input.idea),
            ("Дедлайн", _format_date(report.source_input.deadline)),
            ("Бюджет", _format_money(report.source_input.budget)),
            ("География", report.source_input.geography or "не указана"),
            ("Стейкхолдеры", report.source_input.stakeholders or "не указаны"),
            ("Текущие ресурсы", report.source_input.current_resources or "не указаны"),
            (
                "Технологические ограничения",
                report.source_input.technology_constraints or "не указаны",
            ),
            ("Акценты проекта", report.source_input.project_accents or "не указаны"),
        ]
        _add_table(document, ("Поле", "Значение"), source_rows)

        _add_bullets(document, "Предупреждения", report.warnings)
        _add_bullets(document, "Допущения", report.assumptions)

        _add_heading(document, "Паспорт проекта", level=1)
        _add_paragraph(document, report.passport.goal)
        _add_table(
            document,
            ("Поле", "Значение"),
            [
                ("Целевая аудитория", report.passport.target_audience),
                ("Актуальность для Уральского банка", report.passport.relevance_for_ural_bank),
            ],
        )
        _add_bullets(document, "Задачи", report.passport.tasks)
        _add_bullets(document, "Критерии успеха", report.passport.success_criteria)
        _add_bullets(document, "Риски паспорта проекта", report.passport.risks)
        _add_bullets(document, "Допущения паспорта проекта", report.passport.assumptions)

        _add_heading(document, "Дорожная карта", level=1)
        _add_table(
            document,
            ("Фаза", "Старт", "Финиш", "Контрольные точки"),
            [
                (
                    phase.name,
                    _format_date(phase.start_date),
                    _format_date(phase.end_date),
                    "\n".join(
                        f"{item.title} ({_format_date(item.due_date)}): {item.description}"
                        for item in phase.milestones
                    ),
                )
                for phase in report.roadmap
            ],
        )

        gantt_rows = report.gantt if is_gantt_valid(report.gantt) else build_gantt_from_roadmap(report.roadmap)
        _add_heading(document, "Gantt-like представление", level=1)
        _add_table(
            document,
            ("Фаза", "Период", "Шкала"),
            [(row.phase, row.period, row.timeline) for row in gantt_rows],
        )

        _add_heading(document, "Ресурсы", level=1)
        _add_table(
            document,
            ("Статья", "Сумма", "Комментарий"),
            [
                (item.category, _format_money(item.amount), item.comment)
                for item in report.resources.financial_items
            ],
        )
        _add_paragraph(document, f"Итого: {_format_money(report.resources.financial_total)}")
        _add_bullets(document, "Материально-технические ресурсы", report.resources.material_resources)
        _add_bullets(document, "Информационные ресурсы", report.resources.information_resources)

        _add_heading(document, "Команда проекта", level=1)
        _add_table(
            document,
            ("Роль", "Количество", "Компетенции", "Комментарий"),
            [
                (
                    role.title,
                    str(role.count),
                    ", ".join(role.competencies),
                    role.assignment_comment,
                )
                for role in report.team
            ],
        )

        _add_heading(document, "RACI", level=1)
        _add_table(
            document,
            ("Активность", "R", "A", "C", "I"),
            [
                (
                    item.activity,
                    item.responsible,
                    item.accountable,
                    ", ".join(item.consulted),
                    ", ".join(item.informed),
                )
                for item in report.raci
            ],
        )

        _add_heading(document, "Три концепции", level=1)
        for concept in report.concepts:
            _add_heading(document, concept.name, level=2)
            _add_paragraph(document, concept.key_idea)
            _add_paragraph(document, f"Оценка стоимости: {_format_money(concept.estimated_cost)}")
            _add_paragraph(document, f"Трудоёмкость: {concept.effort_level}")
            _add_paragraph(document, f"Отличия: {concept.differences}")
            _add_bullets(document, "Сценарий", concept.scenario_steps)
            _add_bullets(document, "Преимущества", concept.advantages)
            _add_bullets(document, "Недостатки", concept.disadvantages)
            _add_bullets(document, "Факторы трудоёмкости", concept.effort_factors)

        _add_heading(document, "Рекомендованная концепция", level=1)
        _add_paragraph(document, report.recommended_concept.concept_name)
        _add_paragraph(document, report.recommended_concept.rationale)
        _add_bullets(document, "Риски выбранной концепции", report.recommended_concept.risks)

        if report.presentation_outline:
            _add_heading(document, "Outline презентации", level=1)
            for slide in report.presentation_outline:
                _add_bullets(document, slide.title, slide.bullets)
        if report.defense_script:
            _add_heading(document, "Сценарий защиты", level=1)
            _add_paragraph(document, report.defense_script)

        document.save(path)
    except Exception:
        logger.warning(
            "Failed to export Project Planner DOCX with python-docx; writing minimal fallback.",
            exc_info=True,
        )
        _write_minimal_docx(report, path)
    return path
